"""AFS Output Generator — assembles sections into PDF (HTML), DOCX, and iXBRL."""

from __future__ import annotations

import html
import io
from datetime import datetime
from typing import Any

# ---------------------------------------------------------------------------
# Shared HTML builder
# ---------------------------------------------------------------------------

def _md_table_to_html(md: str) -> str:
    """Convert a markdown table to an HTML table."""
    lines = [ln.strip() for ln in md.strip().splitlines() if ln.strip()]
    if len(lines) < 2:
        return f"<pre>{md}</pre>"

    def _cells(line: str) -> list[str]:
        parts = line.split("|")
        return [c.strip() for c in parts if c.strip() or c.strip() == ""]

    header_cells = _cells(lines[0])
    # Skip separator line (line[1] is usually dashes)
    body_lines = lines[2:] if len(lines) > 2 else []

    html = '<table class="afs-table">\n<thead><tr>'
    for cell in header_cells:
        if cell:
            html += f"<th>{cell}</th>"
    html += "</tr></thead>\n<tbody>\n"
    for line in body_lines:
        cells = _cells(line)
        html += "<tr>"
        for cell in cells:
            if cell or cell == "":
                html += f"<td>{cell}</td>"
        html += "</tr>\n"
    html += "</tbody>\n</table>"
    return html


def _build_html(
    entity_name: str,
    period_start: str,
    period_end: str,
    framework_name: str,
    sections: list[dict[str, Any]],
    *,
    include_xbrl: bool = False,
    standard: str = "ifrs",
) -> str:
    """Build a full HTML document from section content_json structures."""
    safe_entity = html.escape(entity_name)
    safe_period_start = html.escape(period_start)
    safe_period_end = html.escape(period_end)
    safe_framework = html.escape(framework_name)

    xbrl_ns = ""
    xbrl_header = ""
    if include_xbrl:
        taxonomy = "http://xbrl.ifrs.org/taxonomy/2024-03-28/ifrs-full" if "ifrs" in standard.lower() else "http://fasb.org/us-gaap/2024"
        xbrl_ns = ' xmlns:ix="http://www.xbrl.org/2013/inlineXBRL" xmlns:xbrli="http://www.xbrl.org/2003/instance" xmlns:ifrs-full="http://xbrl.ifrs.org/taxonomy/2024-03-28/ifrs-full"'
        xbrl_header = f"""
    <ix:header>
      <ix:references>
        <link:schemaRef xlink:href="{taxonomy}"/>
      </ix:references>
      <ix:resources>
        <xbrli:context id="ctx-current">
          <xbrli:entity><xbrli:identifier scheme="http://www.example.com">{safe_entity}</xbrli:identifier></xbrli:entity>
          <xbrli:period><xbrli:startDate>{safe_period_start}</xbrli:startDate><xbrli:endDate>{safe_period_end}</xbrli:endDate></xbrli:period>
        </xbrli:context>
      </ix:resources>
    </ix:header>"""

    css = """
    body { font-family: 'Times New Roman', serif; margin: 40px; color: #1a1a1a; line-height: 1.6; }
    .cover { text-align: center; padding: 120px 0 80px; page-break-after: always; }
    .cover h1 { font-size: 28px; margin-bottom: 12px; }
    .cover .entity { font-size: 22px; color: #333; }
    .cover .period { font-size: 16px; color: #666; margin-top: 8px; }
    .cover .framework { font-size: 14px; color: #888; margin-top: 24px; }
    .toc { page-break-after: always; }
    .toc h2 { font-size: 20px; border-bottom: 2px solid #333; padding-bottom: 4px; }
    .toc ul { list-style: none; padding: 0; }
    .toc li { padding: 4px 0; font-size: 14px; }
    .section { page-break-before: always; margin-bottom: 32px; }
    .section:first-of-type { page-break-before: avoid; }
    .section h2 { font-size: 18px; border-bottom: 1px solid #ccc; padding-bottom: 4px; margin-bottom: 12px; }
    .section h3 { font-size: 16px; margin: 16px 0 8px; }
    .section p { font-size: 13px; margin: 8px 0; }
    .afs-table { width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 12px; }
    .afs-table th, .afs-table td { border: 1px solid #ccc; padding: 6px 8px; text-align: left; }
    .afs-table th { background: #f5f5f5; font-weight: bold; }
    .references { margin-top: 8px; font-size: 11px; color: #666; }
    .warnings { margin-top: 8px; padding: 8px; background: #fff8e1; border: 1px solid #ffcc02; font-size: 12px; }
    @media print { .section { page-break-before: always; } }
    """

    parts: list[str] = [
        f'<!DOCTYPE html>\n<html{xbrl_ns}>\n<head>',
        f'<meta charset="utf-8"><title>{safe_entity} — Annual Financial Statements</title>',
        f"<style>{css}</style></head>\n<body>",
    ]

    if xbrl_header:
        parts.append(xbrl_header)

    # Cover page
    parts.append(f"""
    <div class="cover">
      <h1>Annual Financial Statements</h1>
      <div class="entity">{safe_entity}</div>
      <div class="period">For the period {safe_period_start} to {safe_period_end}</div>
      <div class="framework">Prepared in accordance with {safe_framework}</div>
      <div class="framework">Generated {datetime.utcnow().strftime('%d %B %Y')}</div>
    </div>""")

    # Table of contents
    parts.append('<div class="toc"><h2>Table of Contents</h2><ul>')
    for i, sec in enumerate(sections, 1):
        title = html.escape(sec.get("title", f"Section {i}"))
        parts.append(f"<li>{i}. {title}</li>")
    parts.append("</ul></div>")

    # Sections
    for i, sec in enumerate(sections, 1):
        content = sec.get("content_json") or {}
        title = html.escape(content.get("title", sec.get("title", f"Section {i}")))
        paragraphs = content.get("paragraphs", [])
        references = content.get("references", [])
        warnings = content.get("warnings", [])

        parts.append(f'<div class="section"><h2>{i}. {title}</h2>')

        for para in paragraphs:
            ptype = para.get("type", "text")
            pcontent = para.get("content", "")
            if ptype == "heading":
                parts.append(f"<h3>{html.escape(pcontent)}</h3>")
            elif ptype == "table":
                parts.append(_md_table_to_html(pcontent))
            else:
                parts.append(f"<p>{html.escape(pcontent)}</p>")

        if references:
            refs_str = ", ".join(html.escape(r) for r in references)
            parts.append(f'<div class="references">References: {refs_str}</div>')

        if warnings:
            parts.append('<div class="warnings"><strong>Notes:</strong><ul>')
            for w in warnings:
                parts.append(f"<li>{html.escape(w)}</li>")
            parts.append("</ul></div>")

        parts.append("</div>")

    parts.append("</body></html>")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# PDF (HTML for browser print)
# ---------------------------------------------------------------------------

def generate_pdf_html(
    entity_name: str,
    period_start: str,
    period_end: str,
    framework_name: str,
    sections: list[dict[str, Any]],
) -> bytes:
    """Generate print-ready HTML for PDF output."""
    html = _build_html(entity_name, period_start, period_end, framework_name, sections)
    return html.encode("utf-8")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def generate_docx(
    entity_name: str,
    period_start: str,
    period_end: str,
    framework_name: str,
    sections: list[dict[str, Any]],
) -> bytes:
    """Generate DOCX from sections using python-docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    # Title page
    title = doc.add_heading("Annual Financial Statements", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    entity_para = doc.add_paragraph()
    entity_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = entity_para.add_run(entity_name)
    run.font.size = Pt(18)

    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period_para.add_run(f"For the period {period_start} to {period_end}")
    run.font.size = Pt(12)

    fw_para = doc.add_paragraph()
    fw_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fw_para.add_run(f"Prepared in accordance with {framework_name}")
    run.font.size = Pt(10)

    doc.add_page_break()

    # Table of contents placeholder
    doc.add_heading("Table of Contents", level=1)
    for i, sec in enumerate(sections, 1):
        title_text = sec.get("title", f"Section {i}")
        doc.add_paragraph(f"{i}. {title_text}", style="List Number")
    doc.add_page_break()

    # Sections
    for i, sec in enumerate(sections, 1):
        content = sec.get("content_json") or {}
        title_text = content.get("title", sec.get("title", f"Section {i}"))
        paragraphs = content.get("paragraphs", [])
        references = content.get("references", [])
        warnings = content.get("warnings", [])

        doc.add_heading(f"{i}. {title_text}", level=1)

        for para in paragraphs:
            ptype = para.get("type", "text")
            pcontent = para.get("content", "")
            if ptype == "heading":
                doc.add_heading(pcontent, level=2)
            elif ptype == "table":
                _add_md_table_to_docx(doc, pcontent)
            else:
                doc.add_paragraph(pcontent)

        if references:
            refs_para = doc.add_paragraph()
            refs_para.add_run("References: ").bold = True
            refs_para.add_run(", ".join(references))

        if warnings:
            warn_para = doc.add_paragraph()
            warn_para.add_run("Notes: ").bold = True
            for w in warnings:
                doc.add_paragraph(f"• {w}")

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_md_table_to_docx(doc: Any, md: str) -> None:
    """Parse a markdown table and add it to the DOCX document."""
    from docx.shared import Pt

    lines = [ln.strip() for ln in md.strip().splitlines() if ln.strip()]
    if len(lines) < 2:
        doc.add_paragraph(md)
        return

    def _cells(line: str) -> list[str]:
        parts = line.split("|")
        return [c.strip() for c in parts if c.strip() or c.strip() == ""]

    header_cells = [c for c in _cells(lines[0]) if c]
    body_lines = lines[2:] if len(lines) > 2 else []
    cols = len(header_cells) or 1

    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"

    # Header
    for j, cell_text in enumerate(header_cells):
        if j < cols:
            cell = table.rows[0].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.bold = True

    # Body rows
    for line in body_lines:
        cells = [c for c in _cells(line) if c or c == ""]
        row = table.add_row()
        for j, cell_text in enumerate(cells):
            if j < cols:
                row.cells[j].text = cell_text


# ---------------------------------------------------------------------------
# iXBRL (Inline XBRL HTML)
# ---------------------------------------------------------------------------

def generate_ixbrl(
    entity_name: str,
    period_start: str,
    period_end: str,
    framework_name: str,
    standard: str,
    sections: list[dict[str, Any]],
) -> bytes:
    """Generate inline XBRL HTML from sections."""
    html = _build_html(
        entity_name, period_start, period_end, framework_name, sections,
        include_xbrl=True, standard=standard,
    )
    return html.encode("utf-8")
